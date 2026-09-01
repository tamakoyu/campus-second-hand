# -*- coding: utf-8 -*-
"""生成《AI 智能校园二手交易平台立项报告》.docx —— 含真实 Word 表格与 AI 版演示图"""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '_py_deps'))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SONG = u'宋体'
HEI = u'黑体'

def set_run_font(run, name=SONG, size=12, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.rPr
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, name=SONG, size=12, bold=False, align=None, indent=True,
             space_after=6, line=1.4):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if indent:
        pf.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, name=name, size=size, bold=bold)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 8)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.3
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, name=HEI, size=16, bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, name=HEI, size=14, bold=True)
    else:
        run = p.add_run(text)
        set_run_font(run, name=HEI, size=12, bold=True)
    return p

def shade_cell(cell, fill='D9E2F3'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def fill_cell(cell, text, bold=False, size=10.5, name=SONG, center=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, name=name, size=size, bold=bold)

def add_table(doc, headers, rows, widths=None, header_fill='D9E2F3'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        fill_cell(hdr[i], h, bold=True, center=True)
        shade_cell(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            fill_cell(cells[i], v)
    if widths:
        t.autofit = False
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), SONG)

# ===== 标题 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run(u'AI 智能校园二手交易平台立项报告')
set_run_font(run, name=HEI, size=22, bold=True)

# ===== 项目基本信息 =====
add_heading(doc, u'项目基本信息', level=2)
add_table(doc,
    [u'项目', u'内容'],
    [
        [u'项目题目', u'AI 智能校园二手物品交易平台'],
        [u'项目类型', u'Web 应用（B/S 前后端分离架构 + AI 服务层）'],
        [u'团队人数', u'5 人'],
        [u'组长', u'陈思瀚'],
        [u'面向用户', u'本校师生'],
        [u'核心技术栈', u'Vue 3 + Vite + Element Plus / Java Spring Boot + MyBatis-Plus / PostgreSQL 18.6 / 大模型 API（智谱 GLM 系列等免费接口）'],
        [u'开发版本规划', u'V0.1（概念验证）→ V0.5（中期检查）→ V1.0（项目验收）'],
    ],
    widths=[4.0, 11.0])

# ===== 一、项目摘要 =====
add_heading(doc, u'一、项目摘要')
add_para(doc, u'本项目聚焦在校大学生校园生活场景，开发一款 AI 赋能的校园二手物品交易 Web 平台，打造"会估价、会描述、有信用、能减碳"的校园闲置循环服务体系。针对校内闲置物品流转困难、定价不透明、交易信任缺失、缺乏绿色导向等现实痛点，平台以 AI 能力为核心创新：AI 智能估价（规则引擎 + 大模型双层估价，解决"卖多少、买多少"的定价难题）、AI 智能发布助手（图片识别自动分类、描述自动生成，发布效率提升至秒级）、AI 智能问答（商品详情页即时答疑）、AI 辅助审核（违禁与虚假信息智能预检）；同时构建校园信任体系（学号实名认证、校园信用分、线上担保交易、交易互评）与绿色低碳机制（交易减碳量可视化、环保积分），形成安全、可信、绿色的校园二手交易闭环。')
add_para(doc, u'平台全部业务流程均在线上完成：发布 → 浏览检索 → AI 沟通 → 担保下单 → 确认收货 → 交易互评，不设置线下驿站、自提点等线下服务，交付方式由买卖双方线上协商确定，平台专注线上交易链路的质量与安全。')
add_para(doc, u'项目采用主流前后端分离开发模式，前端基于 Vue 3 + Vite + Element Plus，后端基于 Java Spring Boot + MyBatis-Plus，数据库采用 PostgreSQL 18.6；AI 能力基于国内免费大模型接口（智谱 GLM-4.7-Flash / GLM-4V-Flash、阿里云百炼等，永久免费或含大额免费额度、国内直连、OpenAI 兼容协议）实现，零 API 成本；AI 调用统一封装并配置规则引擎兜底，保证系统在任何网络环境下均可用。项目完整落地 Web 全栈开发、AI 接口集成、数据库设计、接口联调与团队协同开发流程，具备极高的校园实用性与课程实践价值。')

# ===== 二、项目调研与现状分析 =====
add_heading(doc, u'二、项目调研与现状分析')
add_para(doc, u'为精准定位项目需求、规避现有产品弊端，本团队针对市面主流校园二手交易相关产品开展调研分析，具体调研情况如下：')
add_table(doc,
    [u'对比维度', u'闲鱼平台', u'校园集市公众号', u'校园二手 QQ 群'],
    [
        [u'平台定位', u'全网综合型二手交易平台', u'轻量化校园二手信息发布渠道', u'校园闲置流转渠道（群聊）'],
        [u'功能特点', u'功能体系完善、用户基数庞大', u'支持学生发布、浏览二手物品帖子', u'上手门槛低、即时性强'],
        [u'主要不足', u'面向全网，校园专属交易信息混杂；校外商家、无关交易信息过多；无针对性本校用户筛选机制', u'功能简陋，无完整用户账号权限体系；图片上传与管理能力薄弱；缺失精准检索、分类筛选；后台审核与管理模块缺失', u'消息刷新速度快，历史交易信息无法留存检索；无商品分类与收藏功能；无专人内容审核，虚假信息、诈骗信息频发'],
        [u'交易安全', u'无法保障校园交易的安全性与精准性，不适用封闭式校园场景', u'无法规范化管理校园交易信息', u'交易安全性无法保障'],
        [u'AI 智能能力', u'无 AI 估价、AI 描述生成、AI 问答等能力，定价依赖卖家主观判断', u'无任何 AI 能力，发布依赖手工填写', u'无任何 AI 能力，信息检索依赖人工翻找'],
    ],
    widths=[2.4, 4.2, 4.2, 4.2])
add_para(doc, u'综合调研结果可知，现有二手交易工具要么面向全网、场景不聚焦，要么功能简陋、缺乏系统化管控，且均未将 AI 能力应用于闲置定价、描述生成与智能问答等关键环节，校园 AI 二手交易存在明显创新空白。因此，本团队决定开发一款专属本校学生、AI 赋能、轻量化、可管控、高安全的校园二手物品交易 Web 平台，填补校园 AI 二手交易系统的空白。')

# ===== 三、项目功能需求列表 =====
add_heading(doc, u'三、项目功能需求列表')
add_heading(doc, u'（一）普通学生用户需求', level=3)
add_table(doc,
    [u'编号', u'功能', u'功能说明'],
    [
        [u'F-01', u'账号管理', u'支持学生用户注册、登录与学号实名认证（校验学号、姓名、专业一致性，排除校外人员），自主维护个人基本信息'],
        [u'F-02', u'AI 智能发布', u'上传商品图片后，AI 自动识别商品分类与成色、生成描述草稿并给出建议售价，用户确认或修改后一键上架'],
        [u'F-03', u'商品浏览与检索', u'支持分类筛选、关键词精准搜索与 AI 语义搜索（自然语言理解，如"求购九成新高数教材"），浏览平台全部合规商品'],
        [u'F-04', u'AI 智能推荐', u'基于用户浏览、收藏与成交记录生成"猜你喜欢"个性化推荐位，提升闲置流转效率'],
        [u'F-05', u'商品收藏', u'用户可收藏心仪商品，在个人中心统一查看与管理收藏列表'],
        [u'F-06', u'在线沟通', u'商品详情页 AI 智能问答即时答疑，支持转人工私信，实现线上无障碍沟通'],
        [u'F-07', u'线上担保交易', u'买家下单后货款平台托管（模拟），卖家确认发货、买家确认收货后完成结算，全程线上闭环'],
        [u'F-08', u'交易互评与校园信用分', u'交易完成后双方互评；信用分随成交、好评、履约等行为动态增减，影响发布与交易权限'],
        [u'F-09', u'个人商品与订单管理', u'个人中心管理自己发布的商品（编辑/下架）与全部订单（待付款/待发货/待收货/已完成/已取消）'],
        [u'F-10', u'减碳记录与环保积分', u'每笔交易自动计算避免的碳排放量，生成个人减碳档案并累计环保积分，激励绿色循环行为'],
    ],
    widths=[1.8, 3.2, 10.0])
add_heading(doc, u'（二）平台管理员需求', level=3)
add_table(doc,
    [u'编号', u'功能', u'功能说明'],
    [
        [u'A-01', u'AI 辅助审核', u'AI 对商品标题、描述、图片进行违禁词、敏感图、异常定价预检并给出结论，管理员人工复核后通过或驳回'],
        [u'A-02', u'数据统计与可视化', u'查看用户、商品、交易、减碳总量与趋势，以图表可视化展示热门品类、交易热度、减碳排行等'],
        [u'A-03', u'用户账号与信用管理', u'账号管控、异常账号处理、信用分申诉审核、交易纠纷仲裁'],
    ],
    widths=[1.8, 3.2, 10.0])
add_heading(doc, u'（三）平台创新点概述', level=3)
add_table(doc,
    [u'创新点', u'实现思路', u'用户价值'],
    [
        [u'① AI 智能估价', u'折旧规则引擎兜底 + 大模型结合平台同类成交数据给出建议价区间与理由', u'解决"卖多少钱、买多少钱"的定价难题，买卖信息对称'],
        [u'② AI 智能发布助手', u'多模态大模型识别图片自动分类成色 + 自动生成描述草稿', u'商品发布从"5 分钟手工填写"降至"30 秒确认上架"'],
        [u'③ AI 智能问答', u'商品详情 + 卖家预设问答 + 平台规则拼接上下文，轻量 RAG 即时答疑', u'详情页秒级答疑，减少无效私信，提升沟通效率'],
        [u'④ 校园信任体系', u'学号实名认证 + 校园信用分 + 线上担保交易 + 交易互评', u'从源头过滤校外人员，解决"不敢交易"的信任难题'],
        [u'⑤ 绿色减碳可视化', u'交易减碳量自动计算 + 个人减碳档案 + 环保积分', u'让每次闲置流转都成为可量化的校园环保行动'],
    ],
    widths=[3.2, 6.4, 5.4])

# ===== 四、项目技术方案 =====
add_heading(doc, u'四、项目技术方案')
add_heading(doc, u'（一）系统架构', level=3)
add_para(doc, u'本项目采用成熟的 B/S 前后端分离架构，在传统三层架构之上新增 AI 服务层：所有大模型调用统一封装，规则引擎提供离线兜底，保证 AI 能力可用性与可控性。')
add_table(doc,
    [u'架构层次', u'技术载体', u'职责说明'],
    [
        [u'客户端层', u'浏览器', u'前端访问载体，无需安装客户端'],
        [u'表现层', u'Vue 3 + Vite + Element Plus', u'页面渲染、用户交互，通过 RESTful 接口调用后端'],
        [u'业务层', u'Spring Boot + MyBatis-Plus', u'提供标准化 RESTful 接口，实现业务逻辑与数据库 CRUD'],
        [u'AI 服务层', u'大模型 API（OpenAI 兼容协议）+ 规则引擎', u'AI 估价、描述生成、智能问答、审核预检；统一 AiService 封装，异常自动降级规则引擎'],
        [u'数据层', u'PostgreSQL 18.6 + pgAdmin 4', u'业务数据持久化存储与可视化管理'],
    ],
    widths=[3.0, 5.6, 6.4])
add_heading(doc, u'（二）核心技术栈', level=3)
add_table(doc,
    [u'层级', u'技术选型', u'说明与优势'],
    [
        [u'前端', u'Vue 3 + Vite + Element Plus', u'轻量化、高性能，组件库快速搭建标准化页面，提升开发效率'],
        [u'后端', u'Java Spring Boot + MyBatis-Plus', u'简化配置、快速搭建服务，高效实现数据库 CRUD 操作'],
        [u'数据库', u'PostgreSQL 18.6 + pgAdmin 4', u'开源稳定、安全性高、兼容性强，实现数据库可视化管理'],
        [u'AI 能力', u'智谱 GLM-4.7-Flash（文字）/ GLM-4V-Flash（视觉）等免费大模型接口', u'永久免费、国内直连、OpenAI 兼容；统一 AiService 封装 + 规则引擎兜底 + JSON 结构化输出'],
    ],
    widths=[2.0, 6.0, 7.0])
add_heading(doc, u'（三）AI 能力技术方案（核心创新落地）', level=3)
add_para(doc, u'1. AI 服务统一封装：后端抽象 AiService 统一承接全部大模型调用，采用 OpenAI 兼容协议，供应商可配置切换（智谱、阿里云百炼双通道互备）；API Key 仅存储于后端环境变量，前端不直接接触大模型接口。', indent=False)
add_para(doc, u'2. AI 智能估价（双层架构）：第一层为离线规则引擎，按"品类折旧系数 × 成色系数 × 原价"计算基础估价，保证任何环境可用；第二层由大模型结合平台同类商品近期成交数据给出建议价区间与定价理由，通过 response_format 强制 JSON 结构化输出，解析失败自动降级第一层。', indent=False)
add_para(doc, u'3. AI 智能发布助手：多模态大模型接收商品图片（base64 传参），返回分类、成色与描述草稿的结构化 JSON，自动回填发布表单，卖家确认后上架。', indent=False)
add_para(doc, u'4. AI 智能问答：将商品详情、卖家预设问答、平台交易规则拼接为上下文（轻量 RAG，不引入向量库），实现详情页即时答疑；无法回答时引导转人工私信。', indent=False)
add_para(doc, u'5. 容错与降级：AI 调用统一设置 5~10 秒超时；超时、异常、返回格式错误一律降级到规则引擎或默认文案；双供应商互备，答辩演示前预缓存演示数据，确保演示流程永不中断。', indent=False)
add_heading(doc, u'（四）开发与测试工具', level=3)
add_table(doc,
    [u'工具', u'用途'],
    [
        [u'IDEA', u'后端 Java 代码开发、项目配置与服务调试'],
        [u'VS Code', u'前端 Vue 页面、组件与样式开发'],
        [u'Postman', u'后端接口与 AI 接口联调、参数测试、有效性验证'],
        [u'Git + GitHub', u'代码版本控制、多人同步、分支管理与代码追溯'],
    ],
    widths=[4.0, 11.0])
add_heading(doc, u'（五）运行与管理方式', level=3)
add_para(doc, u'项目前期以本地调试运行为主，完成功能开发与测试；后期可部署至本地服务器，用户通过浏览器直接访问使用。AI 能力依赖的免费大模型接口为云端在线服务，本地调试与部署环境均直接可用。项目全程依托在线文档记录会议内容、跟进开发任务、留存项目资料，保障项目有序推进。')

# ===== 五、团队分工与任务安排 =====
add_heading(doc, u'五、团队分工与任务安排')
add_para(doc, u'本项目团队共 5 人，设置项目组长 1 名、前端开发 2 名、后端开发 2 名（含测试运维），AI 相关模块明确分配到人，保障创新功能落地，具体分工如下：')
add_table(doc,
    [u'成员', u'角色', u'负责模块与职责'],
    [
        [u'陈思瀚', u'组长（后端开发）', u'统筹项目管理，梳理需求、撰写全套立项及项目文档，负责立项答辩；主导用户模块、登录权限模块开发；负责 AiService 统一封装与 AI 智能估价接口；负责代码审查、整体集成与进度把控'],
        [u'范胜洲', u'前端开发', u'负责首页、商品列表、检索、商品详情页面开发与优化；负责 AI 智能问答交互与 AI 推荐位；封装通用公共组件，统一页面样式与交互规范'],
        [u'徐家凯', u'前端开发', u'负责登录注册、个人中心、AI 智能发布表单（图片识别、描述生成、估价展示）、私信页面与管理后台前端页面开发调试'],
        [u'田博', u'后端开发', u'负责数据库表结构设计（含信用分、订单、减碳记录表）；开发商品核心业务接口；实现 AI 估价规则引擎与 AI 自动填表发布链路，对接前端联调'],
        [u'林天楚', u'后端开发（含测试运维）', u'负责私信与 AI 问答接口、后台数据统计接口；实现减碳计算与环保积分接口；编写单元测试、功能与回归测试；负责打包部署与环境维护，整理全套测试文档'],
    ],
    widths=[2.2, 3.0, 9.8])
add_para(doc, u'团队全员均独立维护个人 Git 仓库，定期提交代码、更新开发进度，项目末期统一完成代码合并与整体集成，确保项目完整落地。')

# ===== 六、项目开发计划 =====
add_heading(doc, u'六、项目开发计划（V0.1 ~ V1.0）')
add_table(doc,
    [u'里程碑', u'版本', u'阶段目标', u'主要任务'],
    [
        [u'里程碑 1', u'V0.1（概念验证）', u'验证系统架构与 AI 接入链路可行',
         u'完成 PostgreSQL 数据库表结构整体设计；搭建前后端基础框架；实现登录、注册与学号实名认证；完成 AiService 骨架与大模型 API 联调验证；调通前后端基础接口'],
        [u'里程碑 2', u'V0.5（中期检查）', u'核心业务与主要 AI 功能可用',
         u'实现商品发布、浏览检索、收藏、私信等核心功能；落地 AI 智能估价、AI 描述生成、AI 智能问答；完成全模块前后端联调，修复基础 BUG，优化页面交互'],
        [u'里程碑 3', u'V1.0（项目验收）', u'项目正式交付验收',
         u'实现线上担保交易闭环、校园信用分与互评、AI 辅助审核、数据可视化、减碳可视化与环保积分；完成全功能与兼容性测试、BUG 修复；打包部署；整理立项、开发、测试、使用全套文档'],
    ],
    widths=[2.2, 2.8, 3.6, 6.4])

# ===== 七、项目可行性分析 =====
add_heading(doc, u'七、项目可行性分析')
add_heading(doc, u'（一）技术可行性', level=3)
add_para(doc, u'团队全体成员已系统修读 Web 前端技术、Web 程序设计、数据库应用、软件工程等专业核心课程，熟练掌握 Vue3、Spring Boot、PostgreSQL 等核心技术；AI 能力采用国内免费大模型接口（智谱 GLM 系列永久免费、国内直连，无需付费与境外网络），估价规则引擎为纯业务逻辑，AI 调用统一封装并配置失败降级，技术风险可控。本项目所用技术均为课堂重点教学内容，技术栈成熟、难度适中，工作量适配课程实训学时要求，技术方案科学可行，可顺利完成开发落地。')
add_heading(doc, u'（二）团队可行性', level=3)
add_para(doc, u'团队组织结构完整，分工清晰合理，成员能力互补，AI 模块落实到人。项目开发期间定期召开小组会议，同步开发进度、沟通疑难问题；全程采用 Git + GitHub 进行代码版本管控，支持代码提交、审查、回溯；依托在线文档记录会议纪要、跟进任务进度。组长统筹协调，成员高效协作，可保障项目有序、高效、高质量完成。')
add_heading(doc, u'（三）可行性分析汇总', level=3)
add_table(doc,
    [u'维度', u'分析要点', u'结论'],
    [
        [u'技术可行性', u'核心技术为课堂重点教学内容，技术栈成熟、难度适中；工作量适配实训学时', u'可行'],
        [u'AI 可行性', u'免费大模型 API + 规则引擎兜底 + 统一封装，零 API 成本、失败自动降级', u'可行'],
        [u'团队可行性', u'组织结构完整、分工清晰、AI 模块落实到人；定期例会 + Git 管控 + 在线文档协作', u'可行'],
        [u'资源与运行', u'前期本地调试，后期部署本地服务器；AI 依赖云端免费接口，本地与部署环境均可用', u'可行'],
    ],
    widths=[3.0, 9.6, 2.4])

# ===== 八、系统界面演示 =====
add_heading(doc, u'八、系统界面演示（原型示意图）')
add_para(doc, u'项目尚处于立项开发阶段，以下为按功能需求绘制的系统界面原型示意图（演示用），用于直观展示平台主要页面形态、信息架构与 AI 交互形态。')
_demo_base = os.path.dirname(os.path.abspath(__file__))
_demos = [
    (u'01-平台首页-商品列表.png', u'图 8-1 平台首页（商品列表页，含 AI 估价标记）'),
    (u'02-商品详情页.png', u'图 8-2 商品详情页（含 AI 智能问答）'),
    (u'03-商品发布表单.png', u'图 8-3 AI 智能发布表单页（含 AI 估价与描述生成）'),
    (u'04-个人中心-我的发布.png', u'图 8-4 个人中心（含信用分与减碳记录）'),
    (u'05-管理后台-数据统计.png', u'图 8-5 管理后台（含 AI 辅助审核）'),
]
for _fname, _cap in _demos:
    doc.add_picture(os.path.join(_demo_base, u'演示图-AI', _fname), width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p = doc.add_paragraph()
    _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p.paragraph_format.space_after = Pt(12)
    _run = _p.add_run(_cap)
    set_run_font(_run, name=SONG, size=10.5, bold=True)

# ===== 附：风险分析与应对 =====
add_heading(doc, u'附：风险分析与应对', level=3)
add_table(doc,
    [u'风险类型', u'风险描述', u'潜在影响', u'应对措施'],
    [
        [u'AI 准确性', u'AI 估价偏差、描述或问答生成错误', u'误导用户、影响体验', u'规则引擎兜底 + 卖家人工确认修改 + 数据持续校验'],
        [u'API 稳定性与额度', u'免费接口限流、额度耗尽', u'功能不可用', u'双供应商互备 + 超时降级 + 用量告警与用完即停配置'],
        [u'信息安全', u'违禁、虚假、诈骗商品信息混入平台', u'损害平台可信度与交易安全', u'AI 预检 + 人工审核 + 学号实名认证 + 举报机制'],
        [u'数据安全', u'用户隐私泄露、业务数据丢失', u'造成隐私风险、业务中断', u'权限分级管控 + API Key 仅存后端 + 数据库定期备份'],
        [u'进度风险', u'功能开发延期、前后端联调冲突', u'影响里程碑节点与交付验收', u'周会同步进度 + Git 分支管理 + 预留缓冲时间'],
        [u'技术风险', u'AI 接入需自主查阅资料学习', u'拉长开发周期', u'提前预研 AiService 封装、先行接口联调、成员互助'],
    ],
    widths=[2.0, 4.6, 3.6, 4.8])

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), u'项目立项-AI版.docx')
doc.save(out)
print('saved:', out)
